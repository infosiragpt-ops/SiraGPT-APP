#!/usr/bin/env python3
"""
docx-editor - Targeted DOCX editing skill for SiraGPT agents.

Preserves original structure, styles, and formatting.
Works with any LLM by accepting structured JSON instructions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import sys
from pathlib import Path
from typing import Any, Dict, List

def load_document(path: str) -> Document:
    """Load DOCX document safely."""
    return Document(path)

def find_element_by_target(doc: Document, target: str) -> Any:
    """Find paragraph, table or section by heading text or caption."""
    target_lower = target.lower().strip()
    
    # Search headings
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and target_lower in para.text.lower():
            return para
    
    # Search tables by caption or first cell
    for table in doc.tables:
        if table.rows:
            first_cell = table.rows[0].cells[0].text.lower()
            if target_lower in first_cell:
                return table
    
    # Fallback: search paragraphs
    for para in doc.paragraphs:
        if target_lower in para.text.lower():
            return para
    
    return None

def apply_instruction(doc: Document, instruction: Dict[str, Any]) -> bool:
    """Apply a single edit instruction."""
    op = instruction.get("op")
    target = instruction.get("target")
    
    if op == "fill_table":
        element = find_element_by_target(doc, target)
        if element and hasattr(element, 'rows'):
            rows_data = instruction.get("rows", [])
            # Add rows to table (simplified - real impl would match columns)
            for row_data in rows_data:
                row = element.add_row()
                for i, cell_text in enumerate(row_data):
                    if i < len(row.cells):
                        row.cells[i].text = str(cell_text)
            return True
    
    elif op == "add_section":
        # Insert new content after target
        element = find_element_by_target(doc, target)
        if element:
            # Add heading
            new_heading = doc.add_heading(instruction.get("title", "New Section"), level=2)
            # Add content paragraphs
            for content in instruction.get("content", []):
                p = doc.add_paragraph(content.get("text", ""))
                if content.get("style"):
                    p.style = content["style"]
            return True
    
    elif op == "replace_text":
        element = find_element_by_target(doc, target)
        if element and hasattr(element, 'text'):
            element.text = instruction.get("new_text", "")
            return True
    
    return False

def edit_document(source: str, output: str, instructions: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Main entry point for targeted DOCX editing."""
    doc = load_document(source)
    changes = 0
    
    for instr in instructions:
        if apply_instruction(doc, instr):
            changes += 1
    
    doc.save(output)
    
    return {
        "status": "ok",
        "changes_applied": changes,
        "output": output,
        "source": source
    }

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python docx_editor.py <instructions.json>")
        sys.exit(1)
    
    with open(sys.argv[1]) as f:\n        payload = json.load(f)\n    \n    result = edit_document(\n        payload["source"],
        payload["output"],
        payload["instructions"]
    )
    print(json.dumps(result, indent=2))
