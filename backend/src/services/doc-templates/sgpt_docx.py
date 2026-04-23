"""sgpt_docx — style helpers for siraGPT Word output.

Exposes one-line builders that enforce APA 7 / tesis-UPN / corporate
house style so the LLM-generated snippets stay short and the output
stays polished.

The helpers are opinionated:
  · Times New Roman 12 pt, doble interlineado, sangría 1.27 cm en
    cuerpo (APA 7 es exactamente esto).
  · Márgenes 2.54 cm.
  · Heading levels 1-5 con el tratamiento EXACTO que pide APA 7.
  · Tablas sólo con bordes horizontales (estilo APA).
  · Portada, corrientes superiores (running head), números de página,
    referencias con sangría francesa.

Instruments pre-baked:
  · BAI (Beck Anxiety Inventory) — 21 ítems.
  · WHOQOL-BREF — 26 ítems.
  · PHQ-9, GAD-7 — siempre útiles en tesis psico-sanitarias.

Usage:
    from sgpt_docx import (
        apa_document, apa_cover, apa_heading, apa_paragraph,
        apa_table, apa_references, apa_page_break, apa_table_of_contents,
        instrument_bai, instrument_whoqol_bref,
    )
"""

import io
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Low-level XML helpers ───────────────────────────────────────────────

def _set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # Force east-asian + complex-script fallback to the same font so
    # Greek / accented vowels stay Times New Roman in Word.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def _set_cell_borders(cell, top=False, bottom=False):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for side in ('top', 'left', 'bottom', 'right'):
        existing = tc_borders.find(qn(f'w:{side}'))
        if existing is not None:
            tc_borders.remove(existing)
        el = OxmlElement(f'w:{side}')
        visible = (side == 'top' and top) or (side == 'bottom' and bottom)
        el.set(qn('w:val'), 'single' if visible else 'nil')
        if visible:
            el.set(qn('w:sz'), '8')   # 1.0 pt
            el.set(qn('w:color'), '000000')
        tc_borders.append(el)


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)
    _set_font(run, size=12)


# ─── Public API ─────────────────────────────────────────────────────────

def apa_document(*, font='Times New Roman', size=12, running_head=True):
    """Create a blank docx with APA 7 page setup + default style."""
    doc = Document()

    # Margins 2.54 cm (= 1 inch) on every side.
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Override Normal style to APA defaults.
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(size)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Ensure east-asian fallback matches Normal.
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font); rFonts.set(qn('w:hAnsi'), font); rFonts.set(qn('w:cs'), font)

    if running_head:
        # Header with right-aligned page number, footer empty.
        header = doc.sections[0].header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_page_number(p)

    return doc


def apa_cover(doc, *, title, author, institution, course=None, professor=None, date=None, degree=None):
    """APA 7 student title page. Centred block in the upper third."""
    # 3-4 blank lines to push content down (APA places it in upper middle).
    for _ in range(3):
        p = doc.add_paragraph('')
    # Title — bold, centered, title case.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    _set_font(r, size=12, bold=True)
    doc.add_paragraph('')
    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(author), size=12)
    # Institution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(institution), size=12)
    if degree:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(degree), size=12)
    if course:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(course), size=12)
    if professor:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(professor), size=12)
    if date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(date), size=12)
    apa_page_break(doc)


def apa_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def apa_heading(doc, level, text):
    """APA 7 heading levels 1..5."""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(text), size=12, bold=True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(p.add_run(text), size=12, bold=True)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(p.add_run(text), size=12, bold=True, italic=True)
    elif level == 4:
        # Indented, bold, ends with period — inline with first paragraph.
        p.paragraph_format.left_indent = Cm(1.27)
        _set_font(p.add_run(text.rstrip('.') + '.'), size=12, bold=True)
    elif level == 5:
        p.paragraph_format.left_indent = Cm(1.27)
        _set_font(p.add_run(text.rstrip('.') + '.'), size=12, bold=True, italic=True)
    else:
        _set_font(p.add_run(text), size=12, bold=True)
    return p


def apa_paragraph(doc, text, *, first_line_indent=True, italic=False, bold=False, center=False):
    """Body paragraph with APA defaults (first-line indent 1.27 cm)."""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent and not center:
        p.paragraph_format.first_line_indent = Cm(1.27)
    _set_font(p.add_run(text), size=12, italic=italic, bold=bold)
    return p


def apa_table(doc, *, headers, rows, caption_number=None, caption_title=None, note=None):
    """APA 7 table — only top + bottom + header-separator borders.
    headers: list[str]; rows: list[list[str|number]].
    caption_number like '1' and caption_title render the APA caption block.
    """
    if caption_number and caption_title:
        p = doc.add_paragraph()
        _set_font(p.add_run(f'Tabla {caption_number}'), size=12, bold=True)
        p2 = doc.add_paragraph()
        _set_font(p2.add_run(caption_title), size=12, italic=True)

    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(para.add_run(str(h)), size=11, bold=True)
        _set_cell_borders(cell, top=True, bottom=True)

    # Body rows
    for r_idx, row in enumerate(rows):
        last = (r_idx == len(rows) - 1)
        for i, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[i]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            _set_font(para.add_run(str(val)), size=11)
            _set_cell_borders(cell, top=False, bottom=last)

    if note:
        p = doc.add_paragraph()
        _set_font(p.add_run('Nota. '), size=11, italic=True)
        _set_font(p.add_run(note), size=11)

    return table


def apa_references(doc, references):
    """Hanging-indent references list (APA style).
    references: list[str] already formatted as APA 7.
    """
    apa_heading(doc, 1, 'Referencias')
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)  # hanging
        _set_font(p.add_run(ref), size=12)


def apa_table_of_contents(doc):
    """Stub that inserts Word's built-in TOC field. User will refresh
    the TOC on first open (right-click → update field)."""
    apa_heading(doc, 1, 'Índice')
    p = doc.add_paragraph()
    r = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1); r._r.append(instr); r._r.append(fldChar2); r._r.append(fldChar3)


# ─── Pre-baked psychology instruments ─────────────────────────────────────

BAI_ITEMS_ES = [
    'Torpe o entumecido',
    'Acalorado',
    'Con temblor en las piernas',
    'Incapaz de relajarse',
    'Con temor a que ocurra lo peor',
    'Mareado o que se le va la cabeza',
    'Con latidos del corazón fuertes y acelerados',
    'Inestable',
    'Atemorizado o asustado',
    'Nervioso',
    'Con sensación de bloqueo',
    'Con temblores en las manos',
    'Inquieto, inseguro',
    'Con miedo a perder el control',
    'Con sensación de ahogo',
    'Con temor a morir',
    'Con miedo',
    'Con problemas digestivos',
    'Con desvanecimientos',
    'Con rubor facial',
    'Con sudores, fríos o calientes',
]


def instrument_bai(doc, *, include_cover_text=True):
    """Render the Beck Anxiety Inventory (21 items, 0-3 Likert)."""
    apa_heading(doc, 1, 'Inventario de Ansiedad de Beck (BAI)')
    if include_cover_text:
        apa_paragraph(
            doc,
            'A continuación se presenta una lista de síntomas comunes de '
            'ansiedad. Lea cada ítem cuidadosamente e indique cuánto le '
            'ha molestado cada síntoma durante la ÚLTIMA SEMANA, '
            'incluyendo el día de hoy, marcando con una X la casilla '
            'correspondiente.',
            first_line_indent=False,
        )
        apa_paragraph(
            doc,
            '0 = En absoluto   ·   1 = Levemente   ·   '
            '2 = Moderadamente   ·   3 = Severamente',
            first_line_indent=False, italic=True,
        )
    rows = [[str(i + 1), BAI_ITEMS_ES[i], '☐', '☐', '☐', '☐'] for i in range(len(BAI_ITEMS_ES))]
    apa_table(doc, headers=['#', 'Síntoma', '0', '1', '2', '3'], rows=rows)


WHOQOL_BREF_ES = [
    ('¿Cómo calificaría su calidad de vida?', 'global'),
    ('¿Cuán satisfecho está con su salud?', 'global'),
    ('¿En qué medida piensa que el dolor (físico) le impide hacer lo que necesita?', 'física'),
    ('¿Cuánto necesita de cualquier tratamiento médico para funcionar en su vida diaria?', 'física'),
    ('¿Cuánto disfruta de la vida?', 'psicológica'),
    ('¿En qué medida siente que su vida tiene sentido?', 'psicológica'),
    ('¿Cuál es su capacidad de concentración?', 'psicológica'),
    ('¿Cuánta seguridad siente en su vida diaria?', 'ambiente'),
    ('¿Cuán saludable es el ambiente físico a su alrededor?', 'ambiente'),
    ('¿Tiene energía suficiente para su vida diaria?', 'física'),
    ('¿Es capaz de aceptar su apariencia física?', 'psicológica'),
    ('¿Tiene dinero suficiente para cubrir sus necesidades?', 'ambiente'),
    ('¿Cuán disponible tiene la información que necesita en su vida diaria?', 'ambiente'),
    ('¿En qué medida tiene oportunidad de realizar actividades de ocio?', 'ambiente'),
    ('¿Es capaz de desplazarse de un lugar a otro?', 'física'),
    ('¿Cuán satisfecho está con su sueño?', 'física'),
    ('¿Cuán satisfecho está con su habilidad para realizar actividades de la vida diaria?', 'física'),
    ('¿Cuán satisfecho está con su capacidad de trabajo?', 'física'),
    ('¿Cuán satisfecho está de sí mismo?', 'psicológica'),
    ('¿Cuán satisfecho está con sus relaciones personales?', 'relaciones'),
    ('¿Cuán satisfecho está con su vida sexual?', 'relaciones'),
    ('¿Cuán satisfecho está con el apoyo que obtiene de sus amigos?', 'relaciones'),
    ('¿Cuán satisfecho está con las condiciones del lugar donde vive?', 'ambiente'),
    ('¿Cuán satisfecho está con el acceso que tiene a servicios sanitarios?', 'ambiente'),
    ('¿Cuán satisfecho está con el transporte que utiliza?', 'ambiente'),
    ('¿Con qué frecuencia tiene sentimientos negativos (tristeza, desesperanza, ansiedad)?', 'psicológica'),
]


def instrument_whoqol_bref(doc, *, include_cover_text=True):
    apa_heading(doc, 1, 'WHOQOL-BREF · Calidad de Vida')
    if include_cover_text:
        apa_paragraph(
            doc,
            'Este cuestionario evalúa cómo percibe su calidad de vida, '
            'salud y otras áreas. Por favor, conteste todas las preguntas. '
            'No existen respuestas correctas ni incorrectas.',
            first_line_indent=False,
        )
        apa_paragraph(
            doc,
            '1 = Muy malo/Nada   ·   2 = Poco   ·   3 = Moderado   '
            '·   4 = Bastante   ·   5 = Totalmente/Excelente',
            first_line_indent=False, italic=True,
        )
    rows = [[str(i + 1), q, d, '☐', '☐', '☐', '☐', '☐']
            for i, (q, d) in enumerate(WHOQOL_BREF_ES)]
    apa_table(
        doc,
        headers=['#', 'Pregunta', 'Dominio', '1', '2', '3', '4', '5'],
        rows=rows,
    )


PHQ9_ITEMS = [
    'Poco interés o placer en hacer las cosas',
    'Sentirse decaído, deprimido o sin esperanza',
    'Dificultad para dormir, permanecer dormido o dormir demasiado',
    'Sentirse cansado o con poca energía',
    'Poco apetito o comer en exceso',
    'Sentirse mal consigo mismo, fracasado o que ha defraudado a los demás',
    'Dificultad para concentrarse en cosas como leer o ver televisión',
    'Moverse o hablar tan lento que otros podrían notarlo, o al contrario estar tan inquieto que se ha movido mucho más de lo habitual',
    'Pensar que sería mejor estar muerto o hacerse daño de alguna manera',
]


def instrument_phq9(doc):
    apa_heading(doc, 1, 'PHQ-9 · Cuestionario de Salud del Paciente')
    apa_paragraph(
        doc,
        'Durante las ÚLTIMAS 2 SEMANAS, ¿con qué frecuencia le han '
        'molestado los siguientes problemas?',
        first_line_indent=False,
    )
    apa_paragraph(
        doc,
        '0 = Nunca   ·   1 = Varios días   ·   '
        '2 = Más de la mitad de los días   ·   3 = Casi todos los días',
        first_line_indent=False, italic=True,
    )
    rows = [[str(i + 1), PHQ9_ITEMS[i], '☐', '☐', '☐', '☐'] for i in range(9)]
    apa_table(doc, headers=['#', 'Ítem', '0', '1', '2', '3'], rows=rows)


GAD7_ITEMS = [
    'Sentirse nervioso, ansioso o con los nervios de punta',
    'No poder dejar de preocuparse o no poder controlar la preocupación',
    'Preocuparse demasiado por diferentes cosas',
    'Tener dificultad para relajarse',
    'Estar tan inquieto que no puede quedarse quieto',
    'Enfadarse o irritarse con facilidad',
    'Sentir miedo como si algo terrible pudiera suceder',
]


# ─── docxtpl — data-driven templates for matrices narrativas ──────────────
#
# When the user has a structured data set (rows of references, authors,
# categorías, etc.) and wants them dropped into a styled Word template,
# docxtpl is the cleanest tool. We expose a one-liner that takes a
# template path + context dict and saves the result.
#
# Typical flow the LLM uses:
#     from sgpt_docx import build_narrative_matrix
#     rows = [{"autor":"García (2020)","tema":"X","muestra":"120","hallazgo":"..."}]
#     build_narrative_matrix(OUT_PATH, title="Matriz narrativa", rows=rows,
#                            columns=["autor","tema","muestra","hallazgo"])
#
# No external template file is needed — we build it with python-docx
# using the APA-style helpers and then fill it in-place.

def build_narrative_matrix(out_path, *, title, rows, columns,
                           intro=None, institution=None, author=None):
    """One-shot APA matriz narrativa: portada mini + tabla con las
    columnas elegidas + una nota. Columnas aceptan str keys (misma
    key que los dicts en `rows`)."""
    doc = apa_document()
    if author or institution:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        txt = ' · '.join(x for x in [author, institution] if x)
        _set_font(p.add_run(txt), size=10, italic=True)
    apa_heading(doc, 1, title)
    if intro:
        apa_paragraph(doc, intro)
    # Normalise column keys → labels
    if columns and isinstance(columns[0], tuple):
        keys  = [k for k, _ in columns]
        headers = [lbl for _, lbl in columns]
    else:
        keys = list(columns)
        headers = [k.replace('_', ' ').capitalize() for k in keys]
    body = [[str(r.get(k, '')) for k in keys] for r in rows]
    apa_table(
        doc,
        headers=headers, rows=body,
        caption_number='1',
        caption_title='Matriz narrativa de referencias',
        note=f'N = {len(rows)}. Elaboración propia.',
    )
    doc.save(out_path)
    return out_path


def render_template(template_path, context, out_path):
    """Render a docxtpl Jinja-flavoured .docx template with a context
    dict. The template can use {{ variable }} and {% for %} blocks."""
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    tpl.save(out_path)
    return out_path


def instrument_gad7(doc):
    apa_heading(doc, 1, 'GAD-7 · Trastorno de Ansiedad Generalizada')
    apa_paragraph(
        doc,
        'Durante las ÚLTIMAS 2 SEMANAS, ¿con qué frecuencia le han '
        'molestado los siguientes problemas?',
        first_line_indent=False,
    )
    apa_paragraph(
        doc,
        '0 = Nunca   ·   1 = Varios días   ·   '
        '2 = Más de la mitad de los días   ·   3 = Casi todos los días',
        first_line_indent=False, italic=True,
    )
    rows = [[str(i + 1), GAD7_ITEMS[i], '☐', '☐', '☐', '☐'] for i in range(7)]
    apa_table(doc, headers=['#', 'Ítem', '0', '1', '2', '3'], rows=rows)
